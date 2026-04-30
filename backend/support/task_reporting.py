"""
Görev raporları: yıl/ay, ekip, çalışan, görev detayı ve dışa aktarma verisi.
"""
from __future__ import annotations

from calendar import month_name
from datetime import datetime
import os
from typing import Any

from django.db.models import Count, Q
from django.db.models.functions import Coalesce, TruncMonth
from django.utils import timezone

from accounts.models import Team, User
from .models import Task, TaskTimeEntry, TaskProductionEntry


def _year_bounds(year: int, month: int | None) -> tuple[datetime, datetime]:
    if month is not None and 1 <= month <= 12:
        start = timezone.make_aware(datetime(year, month, 1))
        if month == 12:
            end = timezone.make_aware(datetime(year + 1, 1, 1))
        else:
            end = timezone.make_aware(datetime(year, month + 1, 1))
    else:
        start = timezone.make_aware(datetime(year, 1, 1))
        end = timezone.make_aware(datetime(year + 1, 1, 1))
    return start, end


def base_task_qs(org_id: int):
    return Task.objects.filter(organization_id=org_id).select_related(
        'owner', 'assignee', 'team', 'current_team'
    )


def apply_filters(
    qs,
    *,
    team_id: int | None = None,
    assignee_id: int | None = None,
    status: str | None = None,
):
    if team_id:
        qs = qs.filter(Q(team_id=team_id) | Q(current_team_id=team_id))
    if assignee_id:
        qs = qs.filter(assignee_id=assignee_id)
    if status and status != 'all':
        qs = qs.filter(status=status)
    return qs


def completion_expr():
    """Tamamlanma zamanı: end varsa end, yoksa updated_at (done görevler için)."""
    return Coalesce('end', 'updated_at')


def task_detail_rows(qs) -> list[dict[str, Any]]:
    rows = []
    for t in qs.order_by('-updated_at'):
        team = t.current_team or t.team
        comp = t.end or (t.updated_at if t.status == 'done' else None)
        rows.append(
            {
                'id': t.id,
                'title': t.title,
                'status': t.status,
                'priority': t.priority,
                'mode': t.mode,
                'model_code': t.model_code or '',
                'variant': t.variant or '',
                'quantity': t.quantity,
                'team_name': team.name if team else '',
                'team_id': team.id if team else None,
                'assignee_username': t.assignee.username if t.assignee else '',
                'assignee_id': t.assignee_id,
                'owner_username': t.owner.username if t.owner else '',
                'start': t.start.isoformat() if t.start else '',
                'end': t.end.isoformat() if t.end else '',
                'due': t.due.isoformat() if t.due else '',
                'created_at': t.created_at.isoformat() if t.created_at else '',
                'updated_at': t.updated_at.isoformat() if t.updated_at else '',
                'completed_at': comp.isoformat() if comp else '',
                'planned_hours': float(t.planned_hours or 0),
                'total_planned_minutes': float(t.total_planned_minutes or 0),
                'tags': t.tags or [],
            }
        )
    return rows


def monthly_completed_in_year(org_id: int, year: int, filters: dict) -> list[dict[str, Any]]:
    start, end = _year_bounds(year, None)
    qs = base_task_qs(org_id).filter(status='done')
    qs = apply_filters(qs, **filters)
    qs = qs.annotate(completed_at=completion_expr()).filter(
        completed_at__gte=start,
        completed_at__lt=end,
    )
    agg = (
        qs.annotate(ym=TruncMonth('completed_at'))
        .values('ym')
        .annotate(count=Count('id'))
        .order_by('ym')
    )
    out = []
    for row in agg:
        ym = row['ym']
        if ym is None:
            continue
        key = ym.strftime('%Y-%m')
        out.append({'year_month': key, 'month_label': f"{month_name[ym.month]} {ym.year}", 'completed_count': row['count']})
    return out


def monthly_created_in_year(org_id: int, year: int, filters: dict) -> list[dict[str, Any]]:
    start, end = _year_bounds(year, None)
    qs = base_task_qs(org_id).filter(created_at__gte=start, created_at__lt=end)
    qs = apply_filters(qs, **filters)
    agg = (
        qs.annotate(ym=TruncMonth('created_at'))
        .values('ym')
        .annotate(count=Count('id'))
        .order_by('ym')
    )
    out = []
    for row in agg:
        ym = row['ym']
        if ym is None:
            continue
        out.append({'year_month': ym.strftime('%Y-%m'), 'created_count': row['count']})
    return out


def by_team_summary(org_id: int, year: int, month: int | None, filters: dict) -> list[dict[str, Any]]:
    start, end = _year_bounds(year, month)
    qs_all = apply_filters(base_task_qs(org_id), **filters)
    qs_done = (
        qs_all.filter(status='done')
        .annotate(completed_at=completion_expr())
        .filter(completed_at__gte=start, completed_at__lt=end)
    )
    teams = Team.objects.filter(organization_id=org_id)
    result = []
    for tm in teams.order_by('name'):
        # Tamamlanan: o ay/yılda tamamlanmış ve bu ekibe team veya current_team ile bağlı
        done_n = qs_done.filter(Q(team_id=tm.id) | Q(current_team_id=tm.id)).count()
        active_n = qs_all.filter(~Q(status='done')).filter(Q(team_id=tm.id) | Q(current_team_id=tm.id)).count()
        result.append(
            {
                'team_id': tm.id,
                'team_name': tm.name,
                'tasks_completed_in_period': done_n,
                'tasks_active': active_n,
            }
        )
    return sorted(result, key=lambda x: -x['tasks_completed_in_period'])


def by_user_summary(org_id: int, year: int, month: int | None, filters: dict) -> list[dict[str, Any]]:
    start, end = _year_bounds(year, month)
    base = apply_filters(base_task_qs(org_id), **filters)
    qs_done = (
        base.filter(status='done', assignee__isnull=False)
        .annotate(completed_at=completion_expr())
        .filter(completed_at__gte=start, completed_at__lt=end)
    )
    done_counts = (
        qs_done.values('assignee_id', 'assignee__username')
        .annotate(tasks_completed=Count('id'))
        .order_by('-tasks_completed')
    )
    active_counts = (
        base.filter(~Q(status='done'), assignee__isnull=False)
        .values('assignee_id', 'assignee__username')
        .annotate(tasks_active=Count('id'))
    )
    active_map = {r['assignee_id']: r['tasks_active'] for r in active_counts}
    # Zaman kayıtları (dönem içinde başlayan)
    te_qs = TaskTimeEntry.objects.filter(
        task__organization_id=org_id,
        started_at__gte=start,
        started_at__lt=end,
    )
    if filters.get('team_id'):
        tid = filters['team_id']
        te_qs = te_qs.filter(Q(task__team_id=tid) | Q(task__current_team_id=tid))
    hours_map: dict[int, float] = {}
    for te in te_qs.select_related('user'):
        if not te.user_id:
            continue
        s = te.started_at
        e = te.ended_at or timezone.now()
        hrs = max(0, (e - s).total_seconds() / 3600)
        hours_map[te.user_id] = hours_map.get(te.user_id, 0) + hrs

    rows = []
    for row in done_counts:
        uid = row['assignee_id']
        rows.append(
            {
                'user_id': uid,
                'username': row['assignee__username'] or '',
                'tasks_completed': row['tasks_completed'],
                'tasks_active': active_map.get(uid, 0),
                'hours_logged': round(hours_map.get(uid, 0), 2),
            }
        )
    seen = {r['user_id'] for r in rows}
    for uid, hrs in hours_map.items():
        if uid not in seen and hrs > 0:
            u = User.objects.filter(id=uid, organization_id=org_id).first()
            rows.append(
                {
                    'user_id': uid,
                    'username': u.username if u else str(uid),
                    'tasks_completed': 0,
                    'tasks_active': active_map.get(uid, 0),
                    'hours_logged': round(hrs, 2),
                }
            )
            seen.add(uid)
    return sorted(rows, key=lambda x: (-x['tasks_completed'], -x['hours_logged']))


def build_full_report(org_id: int, year: int, month: int | None, filters: dict) -> dict[str, Any]:
    f = {k: v for k, v in filters.items() if v is not None}
    start, end = _year_bounds(year, month)
    qs = base_task_qs(org_id)
    qs = apply_filters(qs, **f)
    done_in_period = (
        qs.filter(status='done')
        .annotate(completed_at=completion_expr())
        .filter(completed_at__gte=start, completed_at__lt=end)
    )
    touched = qs.filter(updated_at__gte=start, updated_at__lt=end)
    created = qs.filter(created_at__gte=start, created_at__lt=end)
    detail_ids = (
        set(done_in_period.values_list('id', flat=True))
        | set(touched.values_list('id', flat=True))
        | set(created.values_list('id', flat=True))
    )
    detail_qs = qs.filter(id__in=detail_ids) if detail_ids else qs.none()

    monthly_done = monthly_completed_in_year(org_id, year, f)
    monthly_new = monthly_created_in_year(org_id, year, f)
    merged_months: dict[str, dict] = {}
    for m in monthly_done:
        merged_months.setdefault(m['year_month'], {})['completed_count'] = m['completed_count']
        merged_months[m['year_month']]['month_label'] = m.get('month_label', m['year_month'])
    for m in monthly_new:
        k = m['year_month']
        merged_months.setdefault(k, {})['created_count'] = m['created_count']

    def _month_label(ym_key: str, bucket: dict) -> str:
        if bucket.get('month_label'):
            return bucket['month_label']
        if '-' in ym_key:
            y, mo = ym_key.split('-', 1)
            try:
                return f"{month_name[int(mo)]} {y}"
            except (ValueError, IndexError):
                pass
        return ym_key

    monthly_timeline = sorted(
        [
            {
                'year_month': k,
                'month_label': _month_label(k, v),
                'completed_count': v.get('completed_count', 0),
                'created_count': v.get('created_count', 0),
            }
            for k, v in merged_months.items()
        ],
        key=lambda x: x['year_month'],
    )

    if month is not None:
        key = f'{year}-{int(month):02d}'
        monthly_timeline = [m for m in monthly_timeline if m['year_month'] == key]

    # Master analytics
    tasks_for_master = list(detail_qs)
    task_id_set = {t.id for t in tasks_for_master}
    entries_qs = TaskProductionEntry.objects.filter(task__organization_id=org_id, task_id__in=task_id_set).select_related('task', 'user', 'team')
    prod_by_task_line_team_latest: dict[tuple[int, int, int], tuple[datetime, int]] = {}
    worker_line_perf: dict[tuple[int, str, str, int], dict[str, Any]] = {}
    for e in entries_qs:
        line_idx = int(e.product_line_index or 0)
        team_id = int(e.team_id or 0)
        key = (e.task_id, line_idx, team_id)
        current = prod_by_task_line_team_latest.get(key)
        ts = e.created_at or timezone.now()
        if current is None or ts >= current[0]:
            prod_by_task_line_team_latest[key] = (ts, int(e.quantity or 0))
        wname = e.user.username if e.user else '—'
        tname = e.team.name if e.team else '—'
        wkey = (e.user_id or 0, wname, tname, line_idx)
        row = worker_line_perf.setdefault(
            wkey,
            {
                'user_id': e.user_id,
                'username': wname,
                'team_name': tname,
                'product_line_index': line_idx,
                'entry_count': 0,
                'reported_quantity_sum': 0,
                'tasks_touched': set(),
            },
        )
        row['entry_count'] += 1
        row['reported_quantity_sum'] += int(e.quantity or 0)
        row['tasks_touched'].add(e.task_id)

    task_detail_master: list[dict[str, Any]] = []
    fire_rows: list[dict[str, Any]] = []
    stage_rows: list[dict[str, Any]] = []
    stage_qty_rows: list[dict[str, Any]] = []
    category_breakdown_map: dict[str, dict[str, Any]] = {}
    color_breakdown_map: dict[str, dict[str, Any]] = {}
    size_breakdown_map: dict[str, dict[str, Any]] = {}
    line_detail_rows: list[dict[str, Any]] = []
    team_perf_map: dict[str, dict[str, Any]] = {}

    for t in tasks_for_master:
        lines = list(getattr(t, 'product_lines', None) or [])
        target_total = 0
        for ln in lines:
            try:
                target_total += max(0, int((ln or {}).get('quantity') or 0))
            except (TypeError, ValueError):
                continue
        if target_total <= 0:
            target_total = max(1, int(t.quantity or 1))
        wf_ids = [int(x) for x in (t.workflow_team_ids or []) if str(x).isdigit()]
        stage_state = dict(t.workflow_stage_state or {})
        if wf_ids:
            done_vals = []
            for tid in wf_ids:
                st = dict(stage_state.get(str(tid), {}) or {})
                done_vals.append(max(0, int(st.get('qty_done') or 0)))
            realized = min(done_vals) if done_vals else 0
            # Workflow aşama hedef/gerçekleşen (qty_target/qty_done) detayı
            for tid in wf_ids:
                st = dict(stage_state.get(str(tid), {}) or {})
                team_obj = Team.objects.filter(id=tid, organization_id=org_id).first()
                stage_qty_rows.append(
                    {
                        'task_id': t.id,
                        'task_title': t.title,
                        'team_name': team_obj.name if team_obj else str(tid),
                        'qty_target': int(st.get('qty_target') or 0),
                        'qty_done': int(st.get('qty_done') or 0),
                        'stage_done': bool(st.get('stage_done')),
                        'pending_approval': bool(st.get('pending_approval')),
                    }
                )
        else:
            realized = 0
            for ln in lines:
                try:
                    realized += max(0, int((ln or {}).get('qty_produced') or 0))
                except (TypeError, ValueError):
                    continue
        remaining = max(0, target_total - realized)
        task_detail_master.append(
            {
                'task_id': t.id,
                'title': t.title,
                'status': t.status,
                'team_name': (t.current_team or t.team).name if (t.current_team or t.team) else '',
                'target_total': target_total,
                'realized_total': realized,
                'remaining_total': remaining,
                'planned_hours': float(t.planned_hours or 0),
                'start': t.start.isoformat() if t.start else '',
                'end': t.end.isoformat() if t.end else '',
                'updated_at': t.updated_at.isoformat() if t.updated_at else '',
            }
        )
        for idx, ln in enumerate(lines):
            cat = str((ln or {}).get('model_code') or '').strip() or 'Kategori Yok'
            color = str((ln or {}).get('product_color') or '').strip() or 'Renk Yok'
            color_code = str((ln or {}).get('product_color_code') or '').strip()
            sizes = list((ln or {}).get('model_sizes') or [])
            try:
                line_target = max(0, int((ln or {}).get('quantity') or 0))
            except (TypeError, ValueError):
                line_target = 0
            # Kalem bazlı gerçekleşen: workflow varsa tüm ekiplerdeki minimum değer (hattı geçen net adet)
            if wf_ids:
                vals = []
                for tid in wf_ids:
                    st_tid = dict(stage_state.get(str(tid), {}) or {})
                    qmap_tid = dict(st_tid.get('qty_done_by_line') or {})
                    try:
                        vals.append(max(0, int(qmap_tid.get(str(idx), 0) or 0)))
                    except (TypeError, ValueError):
                        vals.append(0)
                line_realized = min(vals) if vals else 0
            else:
                try:
                    line_realized = max(0, int((ln or {}).get('qty_produced') or 0))
                except (TypeError, ValueError):
                    line_realized = 0
            cat_row = category_breakdown_map.setdefault(
                cat,
                {'category_code': cat, 'target_total': 0, 'realized_total': 0},
            )
            cat_row['target_total'] += line_target
            cat_row['realized_total'] += line_realized
            color_key = f'{color} ({color_code})' if color_code else color
            color_row = color_breakdown_map.setdefault(
                color_key,
                {'color': color, 'color_code': color_code, 'target_total': 0, 'realized_total': 0},
            )
            color_row['target_total'] += line_target
            color_row['realized_total'] += line_realized
            if sizes:
                for sz in sizes:
                    sk = str(sz).strip() or 'Olcu Yok'
                    srow = size_breakdown_map.setdefault(
                        sk,
                        {'size': sk, 'target_total': 0, 'realized_total': 0},
                    )
                    srow['target_total'] += line_target
                    srow['realized_total'] += line_realized
            else:
                srow = size_breakdown_map.setdefault(
                    'Olcu Yok',
                    {'size': 'Olcu Yok', 'target_total': 0, 'realized_total': 0},
                )
                srow['target_total'] += line_target
                srow['realized_total'] += line_realized
            line_detail_rows.append(
                {
                    'task_id': t.id,
                    'task_title': t.title,
                    'product_line_index': idx,
                    'category_code': cat,
                    'color': color,
                    'color_code': color_code,
                    'sizes': ', '.join(str(x) for x in sizes) if sizes else '',
                    'target_total': line_target,
                    'realized_total': line_realized,
                    'remaining_total': max(0, line_target - line_realized),
                    'status': t.status,
                }
            )
            fq = float((ln or {}).get('fire_qty') or 0)
            fr = str((ln or {}).get('fire_reason') or '').strip()
            fi = str((ln or {}).get('fire_image_data_url') or '').strip()
            if fq > 0 or fr or fi:
                fire_rows.append(
                    {
                        'task_id': t.id,
                        'title': t.title,
                        'product_line_index': idx,
                        'model_code': str((ln or {}).get('model_code') or ''),
                        'fire_qty': fq,
                        'fire_reason': fr,
                        'fire_where': (t.current_team.name if t.current_team else '') or ((t.team.name if t.team else '') or ''),
                        'fire_image': 'Var' if fi else 'Yok',
                    }
                )
        history = list(t.handover_history or [])
        prev_at = t.start or t.created_at
        for h in history:
            at_raw = h.get('at')
            try:
                at_dt = datetime.fromisoformat(str(at_raw).replace('Z', '+00:00')) if at_raw else None
            except Exception:
                at_dt = None
            if at_dt is None:
                continue
            from_name = h.get('from_team_name') or '—'
            to_name = h.get('to_team_name') or '—'
            dur_h = round(max(0, (at_dt - prev_at).total_seconds()) / 3600, 2) if prev_at else 0
            stage_rows.append(
                {
                    'task_id': t.id,
                    'title': t.title,
                    'from_team': from_name,
                    'to_team': to_name,
                    'at': at_dt.isoformat(),
                    'duration_hours': dur_h,
                    'type': h.get('type') or '',
                    'note': h.get('note') or '',
                }
            )
            perf = team_perf_map.setdefault(
                to_name,
                {'team_name': to_name, 'transition_count': 0, 'total_duration_hours': 0.0, 'max_duration_hours': 0.0},
            )
            perf['transition_count'] += 1
            perf['total_duration_hours'] += dur_h
            perf['max_duration_hours'] = max(perf['max_duration_hours'], dur_h)
            prev_at = at_dt

    worker_rows = []
    for row in worker_line_perf.values():
        cnt = int(row['entry_count'])
        worker_rows.append(
            {
                'user_id': row['user_id'],
                'username': row['username'],
                'team_name': row['team_name'],
                'product_line_index': row['product_line_index'],
                'entry_count': cnt,
                'reported_quantity_sum': row['reported_quantity_sum'],
                'task_count': len(row['tasks_touched']),
                'avg_reported_qty': round((row['reported_quantity_sum'] / cnt), 2) if cnt > 0 else 0,
            }
        )
    worker_rows.sort(key=lambda x: (-x['reported_quantity_sum'], -x['entry_count']))
    team_perf_rows = []
    for v in team_perf_map.values():
        c = int(v['transition_count'])
        team_perf_rows.append(
            {
                **v,
                'avg_duration_hours': round((float(v['total_duration_hours']) / c), 2) if c > 0 else 0,
            }
        )
    team_perf_rows.sort(key=lambda x: -x['avg_duration_hours'])
    category_rows = []
    for v in category_breakdown_map.values():
        tgt = int(v.get('target_total') or 0)
        done = int(v.get('realized_total') or 0)
        category_rows.append(
            {
                'category_code': v.get('category_code') or 'Kategori Yok',
                'target_total': tgt,
                'realized_total': done,
                'remaining_total': max(0, tgt - done),
            }
        )
    category_rows.sort(key=lambda x: x['category_code'])
    color_rows = []
    for v in color_breakdown_map.values():
        tgt = int(v.get('target_total') or 0)
        done = int(v.get('realized_total') or 0)
        color_rows.append(
            {
                'color': v.get('color') or 'Renk Yok',
                'color_code': v.get('color_code') or '',
                'target_total': tgt,
                'realized_total': done,
                'remaining_total': max(0, tgt - done),
            }
        )
    color_rows.sort(key=lambda x: (x['color'], x['color_code']))
    size_rows = []
    for v in size_breakdown_map.values():
        tgt = int(v.get('target_total') or 0)
        done = int(v.get('realized_total') or 0)
        size_rows.append(
            {
                'size': v.get('size') or 'Olcu Yok',
                'target_total': tgt,
                'realized_total': done,
                'remaining_total': max(0, tgt - done),
            }
        )
    size_rows.sort(key=lambda x: x['size'])

    return {
        'year': year,
        'month': month,
        'period_start': start.isoformat(),
        'period_end': end.isoformat(),
        'filters_applied': f,
        'monthly_timeline': monthly_timeline,
        'by_team': by_team_summary(org_id, year, month, f),
        'by_user': by_user_summary(org_id, year, month, f),
        'tasks': task_detail_rows(detail_qs),
        'master': {
            'task_detail': task_detail_master,
            'team_performance': team_perf_rows,
            'worker_line_performance': worker_rows,
            'fire_analysis': fire_rows,
            'stage_durations': sorted(stage_rows, key=lambda r: (r['task_id'], r['at'])),
            'stage_qty_detail': stage_qty_rows,
            'category_breakdown': category_rows,
            'color_breakdown': color_rows,
            'size_breakdown': size_rows,
            'line_detail': line_detail_rows,
        },
    }


def export_xlsx_bytes(data: dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    master = data.get('master') or {}
    by_user = data.get('by_user') or []
    hours_by_user: dict[str, float] = {}
    for r in by_user:
        if r.get('user_id') is None:
            continue
        hours_by_user[str(r.get('user_id'))] = float(r.get('hours_logged') or 0)

    # Özet — aylık
    ws0 = wb.active
    ws0.title = 'Özet'
    ws0.append(['Ay', 'Tamamlanan görev', 'Oluşturulan görev'])
    for row in data.get('monthly_timeline', []):
        ws0.append([row.get('month_label') or row.get('year_month'), row.get('completed_count', 0), row.get('created_count', 0)])
    ws0.append([])
    ws0.append(['Metri̇k', 'Değer'])
    ws0.append(['Toplam görev satırı', len(master.get('task_detail') or [])])
    ws0.append(['Fire kayıt satırı', len(master.get('fire_analysis') or [])])
    ws0.append(['Aşama geçiş satırı', len(master.get('stage_durations') or [])])
    ws0.append(['Kategori kırılım satırı', len(master.get('category_breakdown') or [])])
    ws0.append(['Renk kırılım satırı', len(master.get('color_breakdown') or [])])
    ws0.append(['Olcu kırılım satırı', len(master.get('size_breakdown') or [])])
    ws0.append(['Kalem detay satırı', len(master.get('line_detail') or [])])
    # Görev Detayı
    ws1 = wb.create_sheet('Görev Detayı')
    ws1.append(
        [
            'Görev ID',
            'Başlık',
            'Durum',
            'Aktif Ekip',
            'Toplam Hedef',
            'Gerçekleşen',
            'Kalan',
            'Plan Saat',
            'Başlangıç',
            'Bitiş',
            'Güncellenme',
        ]
    )
    for r in master.get('task_detail', []):
        ws1.append(
            [
                r['task_id'],
                r['title'],
                r['status'],
                r['team_name'],
                r['target_total'],
                r['realized_total'],
                r['remaining_total'],
                r['planned_hours'],
                r['start'],
                r['end'],
                r['updated_at'],
            ]
        )
    # Ekip Performansı
    ws2 = wb.create_sheet('Ekip Performansı')
    ws2.append(['Ekip', 'Geçiş Sayısı', 'Toplam Süre (saat)', 'Ort. Süre (saat)', 'Maks. Süre (saat)'])
    for r in master.get('team_performance', []):
        ws2.append([r['team_name'], r['transition_count'], r['total_duration_hours'], r['avg_duration_hours'], r['max_duration_hours']])
    # Çalışan Performansı
    ws3 = wb.create_sheet('Çalışan Performansı')
    ws3.append(
        [
            'Kullanıcı ID',
            'Kullanıcı',
            'Ekip',
            'Kalem Index',
            'Giriş Sayısı',
            'Bildirilen Toplam',
            'Görev Sayısı',
            'Ortalama',
            'Çalıştığı Saat (toplam)',
        ]
    )
    for r in master.get('worker_line_performance', []):
        ws3.append(
            [
                r['user_id'],
                r['username'],
                r['team_name'],
                r['product_line_index'],
                r['entry_count'],
                r['reported_quantity_sum'],
                r['task_count'],
                r['avg_reported_qty'],
                hours_by_user.get(str(r['user_id']), 0),
            ]
        )
    # Fire Analizi
    ws4 = wb.create_sheet('Fire Analizi')
    ws4.append(['Görev ID', 'Başlık', 'Kalem Index', 'Model Kod', 'Fire Adet', 'Fire Sebep', 'Fire Nerede', 'Fire Görsel'])
    for r in master.get('fire_analysis', []):
        ws4.append(
            [
                r['task_id'],
                r['title'],
                r['product_line_index'],
                r['model_code'],
                r['fire_qty'],
                r['fire_reason'],
                r.get('fire_where', ''),
                r['fire_image'],
            ]
        )
    # Aşama Süreleri
    ws5 = wb.create_sheet('Aşama Süreleri')
    ws5.append(['Görev ID', 'Başlık', 'Önceki Ekip', 'Sonraki Ekip', 'Geçiş Zamanı', 'Aşama Süresi (saat)', 'Tip', 'Not'])
    for r in master.get('stage_durations', []):
        ws5.append([r['task_id'], r['title'], r['from_team'], r['to_team'], r['at'], r['duration_hours'], r['type'], r['note']])

    # Aşama Qty Detayı
    ws6 = wb.create_sheet('Aşama Qty Detayı')
    ws6.append(['Görev ID', 'Görev Başlık', 'Ekip', 'Hedef', 'Gerçekleşen', 'Aşama Done', 'Pending Approval'])
    for r in master.get('stage_qty_detail', []):
        ws6.append(
            [
                r['task_id'],
                r['task_title'],
                r['team_name'],
                r['qty_target'],
                r['qty_done'],
                'Var' if r['stage_done'] else 'Yok',
                'Var' if r['pending_approval'] else 'Yok',
            ]
        )

    # Kategori kırılımı (Kasa/Pervaz/Çıta/Panel vb.)
    ws7 = wb.create_sheet('Kategori Kırılımı')
    ws7.append(['Kategori Kodu', 'Toplam Hedef', 'Gerçekleşen', 'Kalan'])
    for r in master.get('category_breakdown', []):
        ws7.append([r['category_code'], r['target_total'], r['realized_total'], r['remaining_total']])
    ws8 = wb.create_sheet('Renk Kırılımı')
    ws8.append(['Renk', 'Renk Kodu', 'Toplam Hedef', 'Gerçekleşen', 'Kalan'])
    for r in master.get('color_breakdown', []):
        ws8.append([r['color'], r.get('color_code', ''), r['target_total'], r['realized_total'], r['remaining_total']])
    ws9 = wb.create_sheet('Olcu Kırılımı')
    ws9.append(['Olcu', 'Toplam Hedef', 'Gerçekleşen', 'Kalan'])
    for r in master.get('size_breakdown', []):
        ws9.append([r['size'], r['target_total'], r['realized_total'], r['remaining_total']])
    ws10 = wb.create_sheet('Kalem Detayı')
    ws10.append(['Görev ID', 'Görev', 'Kalem', 'Kategori', 'Renk', 'Renk Kod', 'Olculer', 'Hedef', 'Gerçekleşen', 'Kalan', 'Durum'])
    for r in master.get('line_detail', []):
        ws10.append(
            [
                r['task_id'],
                r['task_title'],
                r['product_line_index'],
                r['category_code'],
                r['color'],
                r.get('color_code', ''),
                r.get('sizes', ''),
                r['target_total'],
                r['realized_total'],
                r['remaining_total'],
                r['status'],
            ]
        )

    # Legacy detay (geri uyumluluk)
    ws6 = wb.create_sheet('Görev detay (legacy)')
    headers = [
        'ID',
        'Başlık',
        'Durum',
        'Öncelik',
        'Mod',
        'Model',
        'Varyant',
        'Adet',
        'Ekip',
        'Atanan',
        'Sahip',
        'Başlangıç',
        'Bitiş',
        'Vade',
        'Oluşturulma',
        'Güncellenme',
        'Tamamlanma (tahmini)',
        'Plan saat',
        'Plan dk',
    ]
    ws6.append(headers)
    for t in data.get('tasks', []):
        ws6.append(
            [
                t['id'],
                t['title'],
                t['status'],
                t['priority'],
                t['mode'],
                t['model_code'],
                t['variant'],
                t['quantity'],
                t['team_name'],
                t['assignee_username'],
                t['owner_username'],
                t['start'],
                t['end'],
                t['due'],
                t['created_at'],
                t['updated_at'],
                t['completed_at'],
                t['planned_hours'],
                t['total_planned_minutes'],
            ]
        )
    for ws in wb.worksheets:
        if ws.max_row > 0:
            for cell in ws[1]:
                cell.font = Font(bold=True)
    bio = __import__('io')
    buf = bio.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_docx_bytes(data: dict[str, Any], title: str = 'Görev Raporu') -> bytes:
    from docx import Document

    def _grid(tbl):
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass

    doc = Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph()
    p.add_run(f"Yıl: {data.get('year')}")
    if data.get('month'):
        p.add_run(f"  |  Ay: {data['month']}")
    p.add_run(f"\nDönem: {data.get('period_start', '')} — {data.get('period_end', '')}")

    doc.add_heading('Aylık özet', level=1)
    t1 = doc.add_table(rows=1, cols=3)
    _grid(t1)
    t1.rows[0].cells[0].text = 'Ay'
    t1.rows[0].cells[1].text = 'Tamamlanan'
    t1.rows[0].cells[2].text = 'Oluşturulan'
    for row in data.get('monthly_timeline', []):
        r = t1.add_row().cells
        r[0].text = str(row.get('month_label') or row.get('year_month'))
        r[1].text = str(row.get('completed_count', 0))
        r[2].text = str(row.get('created_count', 0))

    doc.add_heading('Ekip bazlı', level=1)
    t2 = doc.add_table(rows=1, cols=3)
    _grid(t2)
    hdr2 = t2.rows[0].cells
    hdr2[0].text = 'Ekip'
    hdr2[1].text = 'Tamamlanan (dönem)'
    hdr2[2].text = 'Aktif'
    for r in data.get('by_team', []):
        row = t2.add_row().cells
        row[0].text = r['team_name']
        row[1].text = str(r['tasks_completed_in_period'])
        row[2].text = str(r['tasks_active'])

    doc.add_heading('Çalışan bazlı', level=1)
    t3 = doc.add_table(rows=1, cols=4)
    _grid(t3)
    hdr3 = t3.rows[0].cells
    hdr3[0].text = 'Kullanıcı'
    hdr3[1].text = 'Tamamlanan'
    hdr3[2].text = 'Aktif'
    hdr3[3].text = 'Saat'
    for r in data.get('by_user', []):
        row = t3.add_row().cells
        row[0].text = r['username']
        row[1].text = str(r['tasks_completed'])
        row[2].text = str(r['tasks_active'])
        row[3].text = str(r['hours_logged'])

    doc.add_heading('Görev listesi (özet)', level=1)
    for t in data.get('tasks', [])[:500]:
        doc.add_paragraph(f"[{t['status']}] {t['title']} — Ekip: {t['team_name']} — Atanan: {t['assignee_username']}", style='List Bullet')

    bio = __import__('io')
    buf = bio.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_cnc_docx_bytes(data: dict[str, Any]) -> bytes:
    """
    Gunluk uretim faaliyet formatinda doldurulmus DOCX uretir.
    Sunucuda sablon dosyasi bulunursa onu baz alir; yoksa ayni duzen koddan kurulur.
    """
    from docx import Document
    def _grid(tbl):
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass

    template_candidates = [
        '/app/docs/2-CNC .docx',
        '/app/belgeler/2-CNC .docx',
    ]
    doc = None
    for p in template_candidates:
        if os.path.exists(p):
            try:
                doc = Document(p)
                break
            except Exception:
                doc = None
    if doc is None:
        doc = Document()
        doc.add_heading('GUNLUK URETIM FAALIYET RAPORU', 0)
        doc.add_paragraph('BOLUM: TUM URETIM EKIPLERI')

    master = data.get('master') or {}
    task_rows = master.get('task_detail') or []
    fire_rows = master.get('fire_analysis') or []
    stage_qty_rows = master.get('stage_qty_detail') or []
    category_rows = master.get('category_breakdown') or []
    line_detail_rows = master.get('line_detail') or []

    doc.add_paragraph(
        f"Donem: {data.get('period_start', '')} - {data.get('period_end', '')}\n"
        f"Toplam gorev: {len(task_rows)}"
    )

    # Ekip bazli ozet: kac ekip calismis, ne kadar uretmis, kac gorevde bulunmus
    team_agg: dict[str, dict[str, Any]] = {}
    for r in stage_qty_rows:
        tname = str(r.get('team_name') or '—')
        row = team_agg.setdefault(
            tname,
            {'team_name': tname, 'produced_total': 0, 'tasks': set()},
        )
        row['produced_total'] += max(0, int(r.get('qty_done') or 0))
        row['tasks'].add(int(r.get('task_id') or 0))
    teams_worked = len(team_agg)
    doc.add_paragraph(f"Calisan ekip sayisi: {teams_worked}")

    doc.add_heading('KATEGORI KIRILIMI', level=1)
    t_cat = doc.add_table(rows=1, cols=4)
    _grid(t_cat)
    hc = t_cat.rows[0].cells
    hc[0].text = 'KATEGORI KODU'
    hc[1].text = 'TOPLAM HEDEF'
    hc[2].text = 'GERCEKLESEN'
    hc[3].text = 'KALAN'
    for r in category_rows[:200]:
        rr = t_cat.add_row().cells
        rr[0].text = str(r.get('category_code') or '')
        rr[1].text = str(r.get('target_total') or 0)
        rr[2].text = str(r.get('realized_total') or 0)
        rr[3].text = str(r.get('remaining_total') or 0)
    doc.add_heading('GOREV ICI KATEGORI OZETI (ILK 20)', level=1)
    t_line = doc.add_table(rows=1, cols=6)
    _grid(t_line)
    hl = t_line.rows[0].cells
    hl[0].text = 'GOREV'
    hl[1].text = 'KATEGORI'
    hl[2].text = 'HEDEF TOPLAM'
    hl[3].text = 'GERCEKLESEN TOPLAM'
    hl[4].text = 'KALAN'
    hl[5].text = 'ADET SATIR SAYISI'
    grouped: dict[tuple[str, str], dict[str, Any]] = {}
    for r in line_detail_rows:
        task_title = str(r.get('task_title') or '')
        cat = str(r.get('category_code') or 'Kategori Yok')
        key = (task_title, cat)
        row = grouped.setdefault(
            key,
            {
                'task_title': task_title,
                'category_code': cat,
                'target_total': 0,
                'realized_total': 0,
                'line_count': 0,
            },
        )
        row['target_total'] += int(r.get('target_total') or 0)
        row['realized_total'] += int(r.get('realized_total') or 0)
        row['line_count'] += 1
    grouped_rows = sorted(grouped.values(), key=lambda x: (x['task_title'], x['category_code']))
    for r in grouped_rows[:20]:
        rr = t_line.add_row().cells
        rr[0].text = str(r.get('task_title') or '')
        rr[1].text = str(r.get('category_code') or '')
        rr[2].text = str(r.get('target_total') or 0)
        rr[3].text = str(r.get('realized_total') or 0)
        rr[4].text = str(max(0, int(r.get('target_total') or 0) - int(r.get('realized_total') or 0)))
        rr[5].text = str(r.get('line_count') or 0)

    # Görev bazlı kategori kırılımı: sayı farklılıklarını izlemek için
    doc.add_heading('GOREV BAZLI KATEGORI DETAYI', level=1)
    t_cat_task = doc.add_table(rows=1, cols=6)
    _grid(t_cat_task)
    hct = t_cat_task.rows[0].cells
    hct[0].text = 'GOREV ID'
    hct[1].text = 'GOREV'
    hct[2].text = 'KATEGORI KODU'
    hct[3].text = 'HEDEF'
    hct[4].text = 'GERCEKLESEN'
    hct[5].text = 'KALAN'
    # task_detail içinden kategori satırlarını türet
    # (excelden gelen model_code bilgisi product line üzerinde olduğu için master hesapta mevcut)
    # not: burada task başına kategori dağılımını doğrudan task_detail+fire/title ile gösteriyoruz.
    # daha ileri detay için xlsx "Kategori Kırılımı" sheeti referans alınabilir.
    for t in (master.get('task_detail') or [])[:500]:
        # Bu tabloda görev satırını tek satır veriyoruz; kategori kırılımı üst tabloda toplu.
        rr = t_cat_task.add_row().cells
        rr[0].text = str(t.get('task_id') or '')
        rr[1].text = str(t.get('title') or '')
        rr[2].text = 'Toplu'
        rr[3].text = str(t.get('target_total') or 0)
        rr[4].text = str(t.get('realized_total') or 0)
        rr[5].text = str(t.get('remaining_total') or 0)

    doc.add_heading('EKIP BAZLI OZET', level=1)
    t_team = doc.add_table(rows=1, cols=4)
    _grid(t_team)
    h = t_team.rows[0].cells
    h[0].text = 'BOLUM'
    h[1].text = 'URETILEN TOPLAM'
    h[2].text = 'BULUNDUGU GOREV SAYISI'
    h[3].text = 'ACIKLAMA'
    for r in sorted(team_agg.values(), key=lambda x: (-x['produced_total'], x['team_name']))[:200]:
        rr = t_team.add_row().cells
        rr[0].text = str(r.get('team_name') or '')
        rr[1].text = str(r.get('produced_total') or 0)
        rr[2].text = str(len(r.get('tasks') or []))
        rr[3].text = 'Gunluk ekip performansi'

    doc.add_heading('SIPARIS / GOREV DETAYI', level=1)
    t_tasks = doc.add_table(rows=1, cols=8)
    _grid(t_tasks)
    h2 = t_tasks.rows[0].cells
    h2[0].text = 'IS EMRI KODU'
    h2[1].text = 'GOREV'
    h2[2].text = 'AKTIF BOLUM'
    h2[3].text = 'SIPARIS SATIR ADET'
    h2[4].text = 'YAPILAN PANEL ADET'
    h2[5].text = 'SIP.KALAN ADET'
    h2[6].text = 'DURUM'
    h2[7].text = 'GUNCELLENME'
    for r in task_rows[:500]:
        rr = t_tasks.add_row().cells
        rr[0].text = str(r.get('task_id') or '')
        rr[1].text = str(r.get('title') or '')
        rr[2].text = str(r.get('team_name') or '')
        rr[3].text = str(r.get('target_total') or 0)
        rr[4].text = str(r.get('realized_total') or 0)
        rr[5].text = str(r.get('remaining_total') or 0)
        rr[6].text = str(r.get('status') or '')
        rr[7].text = str(r.get('updated_at') or '')

    doc.add_heading('FIRE BILGILERI', level=1)
    t_fire = doc.add_table(rows=1, cols=7)
    _grid(t_fire)
    h3 = t_fire.rows[0].cells
    h3[0].text = 'GOREV ID'
    h3[1].text = 'MODEL KODU'
    h3[2].text = 'FIRE ADET'
    h3[3].text = 'FIRE SEBEP'
    h3[4].text = 'FIRE NEREDE'
    h3[5].text = 'FIRE GORSEL'
    h3[6].text = 'ACIKLAMA'
    for r in fire_rows[:500]:
        rr = t_fire.add_row().cells
        rr[0].text = str(r.get('task_id') or '')
        rr[1].text = str(r.get('model_code') or '')
        rr[2].text = str(r.get('fire_qty') or 0)
        rr[3].text = str(r.get('fire_reason') or '')
        rr[4].text = str(r.get('fire_where') or '')
        rr[5].text = str(r.get('fire_image') or '')
        rr[6].text = str(r.get('title') or '')

    bio = __import__('io')
    buf = bio.BytesIO()
    doc.save(buf)
    return buf.getvalue()
